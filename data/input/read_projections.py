#!/usr/bin/env python3
"""
Read the sample and projections document
"""
from docx import Document

doc = Document('sample and projections.docx')

print("=" * 70)
print("SAMPLE AND PROJECTIONS DOCUMENT")
print("=" * 70)
print()

# Print all paragraphs
for para in doc.paragraphs:
    if para.text.strip():
        print(para.text)
        print()

# Print all tables
for i, table in enumerate(doc.tables):
    print(f"\n{'='*70}")
    print(f"TABLE {i+1}")
    print('='*70)
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        print(" | ".join(cells))
    print()
